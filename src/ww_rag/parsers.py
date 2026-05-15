from __future__ import annotations

import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from ww_rag.models import ParsedSection


VIDEO_EXTENSIONS = {".mp4", ".mov", ".m4a"}


def parse_file(path: Path) -> list[ParsedSection]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return parse_text_file(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".pptx":
        return parse_pptx(path)
    if suffix == ".xlsx":
        return parse_xlsx(path)
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".drawio":
        return parse_drawio(path)
    if suffix == ".msg":
        return parse_msg_placeholder(path)
    if suffix in VIDEO_EXTENSIONS:
        return parse_video_transcript(path)
    return []


def parse_text_file(path: Path) -> list[ParsedSection]:
    text = _read_text(path)
    return _sections_from_text(text, "text")


def parse_docx(path: Path) -> list[ParsedSection]:
    try:
        from docx import Document
    except Exception:
        return parse_docx_zip(path)
    doc = Document(str(path))
    sections: list[ParsedSection] = []
    current_title = "document"
    buffer: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name if para.style else "").lower()
        if "heading" in style and buffer:
            sections.append(ParsedSection(text="\n".join(buffer), location=current_title))
            buffer = []
            current_title = text
        elif "heading" in style:
            current_title = text
        else:
            buffer.append(text)
    if buffer:
        sections.append(ParsedSection(text="\n".join(buffer), location=current_title))
    return sections or [ParsedSection(text="\n".join(p.text for p in doc.paragraphs if p.text), location="document")]


def parse_docx_zip(path: Path) -> list[ParsedSection]:
    try:
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml")
    except Exception:
        return [ParsedSection(text=f"DOCX parser dependency is unavailable for {path.name}", location="parser")]
    texts = _xml_text_nodes(xml)
    return _sections_from_text("\n".join(texts), "document")


def parse_pptx(path: Path) -> list[ParsedSection]:
    sections: list[ParsedSection] = []
    try:
        with zipfile.ZipFile(path) as archive:
            slide_names = sorted(
                name for name in archive.namelist() if name.startswith("ppt/slides/slide") and name.endswith(".xml")
            )
            for index, name in enumerate(slide_names, start=1):
                texts = _xml_text_nodes(archive.read(name))
                if texts:
                    sections.append(ParsedSection(text="\n".join(texts), location=f"slide {index}", metadata={"slide": index}))
    except Exception as exc:
        return [ParsedSection(text=f"PPTX parser failed for {path.name}: {exc}", location="parser")]
    return sections


def parse_xlsx(path: Path) -> list[ParsedSection]:
    try:
        import openpyxl
    except Exception:
        return parse_xlsx_zip(path)
    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    sections: list[ParsedSection] = []
    for sheet in workbook.worksheets:
        rows: list[str] = []
        min_row = None
        max_row = 0
        max_col = 0
        for row_index, row in enumerate(sheet.iter_rows(), start=1):
            values = [cell.value for cell in row]
            if all(value is None or str(value).strip() == "" for value in values):
                if rows:
                    sections.extend(_xlsx_block(sheet.title, rows, min_row or 1, max_row, max_col))
                    rows = []
                    min_row = None
                continue
            if min_row is None:
                min_row = row_index
            max_row = row_index
            max_col = max(max_col, len(values))
            rows.append(" | ".join("" if value is None else str(value) for value in values))
        if rows:
            sections.extend(_xlsx_block(sheet.title, rows, min_row or 1, max_row, max_col))
    return sections


def parse_xlsx_zip(path: Path) -> list[ParsedSection]:
    try:
        with zipfile.ZipFile(path) as archive:
            shared_strings = _xlsx_shared_strings(archive)
            sheet_names = sorted(name for name in archive.namelist() if name.startswith("xl/worksheets/sheet") and name.endswith(".xml"))
            sections: list[ParsedSection] = []
            for index, name in enumerate(sheet_names, start=1):
                rows = _xlsx_sheet_rows(archive.read(name), shared_strings)
                if rows:
                    sections.append(
                        ParsedSection(
                            text="\n".join(rows),
                            location=f"sheet {index}",
                            metadata={"sheet": str(index), "cell_range": f"R1:R{len(rows)}"},
                        )
                    )
            return sections
    except Exception:
        return [ParsedSection(text=f"XLSX parser dependency is unavailable for {path.name}", location="parser")]


def parse_pdf(path: Path) -> list[ParsedSection]:
    try:
        from pypdf import PdfReader
    except Exception:
        return [ParsedSection(text=f"PDF parser dependency is unavailable for {path.name}", location="parser")]
    reader = PdfReader(str(path))
    sections: list[ParsedSection] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            sections.append(ParsedSection(text=text, location=f"page {index}", metadata={"page": index}))
    return sections


def parse_video_transcript(path: Path) -> list[ParsedSection]:
    transcript = _find_transcript(path)
    if transcript is None:
        return [
            ParsedSection(
                text=f"No transcript sidecar was found for video/audio file {path.name}. Provide {path.stem}.txt, .vtt, or .srt to index the meeting content.",
                location="transcript missing",
                metadata={"needs_transcription": True},
            )
        ]
    text = _read_text(transcript)
    if transcript.suffix.lower() in {".vtt", ".srt"}:
        return _sections_from_subtitles(text)
    return _sections_from_text(text, "transcript")


def parse_drawio(path: Path) -> list[ParsedSection]:
    text = _read_text(path)
    cleaned = re.sub(r"<[^>]+>", " ", text)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return [ParsedSection(text=cleaned, location="diagram")] if cleaned else []


def parse_msg_placeholder(path: Path) -> list[ParsedSection]:
    return [
        ParsedSection(
            text=f"Outlook MSG file {path.name} was discovered but full .msg parsing is not enabled. Export this email as .txt, .html, or .pdf to index the message body.",
            location="parser",
            metadata={"needs_msg_export": True},
        )
    ]


def _find_transcript(path: Path) -> Path | None:
    for suffix in (".txt", ".vtt", ".srt"):
        candidate = path.with_suffix(suffix)
        if candidate.exists():
            return candidate
    return None


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="ignore")


def _xml_text_nodes(xml_bytes: bytes) -> list[str]:
    root = ET.fromstring(xml_bytes)
    texts: list[str] = []
    for node in root.iter():
        if node.tag.endswith("}t") and node.text:
            value = node.text.strip()
            if value:
                texts.append(value)
    return texts


def _xlsx_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    try:
        root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
    except KeyError:
        return []
    strings: list[str] = []
    for si in root:
        parts = []
        for node in si.iter():
            if node.tag.endswith("}t") and node.text:
                parts.append(node.text)
        strings.append("".join(parts))
    return strings


def _xlsx_sheet_rows(xml_bytes: bytes, shared_strings: list[str]) -> list[str]:
    root = ET.fromstring(xml_bytes)
    rows: list[str] = []
    for row in root.iter():
        if not row.tag.endswith("}row"):
            continue
        values: list[str] = []
        for cell in row:
            if not cell.tag.endswith("}c"):
                continue
            cell_type = cell.attrib.get("t")
            value = ""
            for child in cell:
                if child.tag.endswith("}v") and child.text:
                    value = child.text
                    break
            if cell_type == "s" and value.isdigit():
                idx = int(value)
                value = shared_strings[idx] if idx < len(shared_strings) else value
            if value:
                values.append(value)
        if values:
            rows.append(" | ".join(values))
    return rows


def _sections_from_text(text: str, default_location: str) -> list[ParsedSection]:
    blocks = [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]
    return [ParsedSection(text=block, location=f"{default_location} {index}") for index, block in enumerate(blocks, start=1)]


def _sections_from_subtitles(text: str) -> list[ParsedSection]:
    sections: list[ParsedSection] = []
    current_time = "00:00:00"
    buffer: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line or line.upper() == "WEBVTT" or line.isdigit():
            continue
        if "-->" in line:
            if buffer:
                sections.append(ParsedSection(text=" ".join(buffer), location=f"time {current_time}", metadata={"timestamp": current_time}))
                buffer = []
            current_time = line.split("-->", 1)[0].strip().replace(",", ".")
        else:
            buffer.append(line)
    if buffer:
        sections.append(ParsedSection(text=" ".join(buffer), location=f"time {current_time}", metadata={"timestamp": current_time}))
    return sections


def _xlsx_block(sheet: str, rows: list[str], min_row: int, max_row: int, max_col: int) -> list[ParsedSection]:
    cell_range = f"R{min_row}:R{max_row}"
    text = "\n".join(rows)
    return [ParsedSection(text=text, location=f"sheet {sheet} {cell_range}", metadata={"sheet": sheet, "cell_range": cell_range, "max_col": max_col})]

